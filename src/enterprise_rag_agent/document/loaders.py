from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Any


@dataclass
class ParsedDocument:
    text: str
    filename: str
    metadata: dict[str, Any] = field(default_factory=dict)


class DocumentLoader:
    """按扩展名解析 PDF、Word、Markdown、TXT。"""

    def load(self, path: str | Path) -> ParsedDocument:
        path = Path(path)
        data = path.read_bytes()
        return self.load_bytes(data, path.name)

    def load_bytes(self, data: bytes, filename: str) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            text = self._load_pdf(data)
            kind = "pdf"
        elif suffix in {".docx", ".doc"}:
            text = self._load_docx(data)
            kind = "word"
        elif suffix in {".md", ".markdown", ".txt", ".text"}:
            text = self._load_text(data)
            kind = "markdown" if suffix.startswith(".md") else "text"
        else:
            raise ValueError(f"不支持的文件类型: {suffix}")
        return ParsedDocument(
            text=text,
            filename=filename,
            metadata={"filename": filename, "extension": suffix, "kind": kind},
        )

    def _load_pdf(self, data: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise RuntimeError("解析 PDF 需要安装 pypdf：pip install pypdf") from exc

        import io

        reader = PdfReader(io.BytesIO(data))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n\n".join(pages)

    def _load_docx(self, data: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError("解析 Word 需要安装 python-docx：pip install python-docx") from exc

        import io

        document = Document(io.BytesIO(data))
        parts: list[str] = []
        for paragraph in document.paragraphs:
            if paragraph.text.strip():
                parts.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def _load_text(self, data: bytes) -> str:
        for encoding in ("utf-8", "gb18030", "utf-16"):
            try:
                return data.decode(encoding)
            except UnicodeDecodeError:
                continue
        return data.decode("utf-8", errors="ignore")

